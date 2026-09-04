"""Genera la propuesta tecnica en DOCX.

QUE PASABA ANTES: NO SE GENERABA NUNCA

  Este modulo no tenia importadores. Y `zip_builder` incluye
  `propuesta_tecnica_{id}.docx` "si existe", asi que el expediente se armaba
  sin ella y sin decir nada. En un concurso publico o una licitacion publica la
  propuesta tecnica es el documento que puntua: un expediente sin ella no es un
  expediente con una carpeta de menos, es una oferta que no compite.

  Ahora la crea la ruta que el usuario ya pulsaba ("Generar documentos"), en el
  mismo sitio donde `zip_builder` la busca.
"""
import logging
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from shared import ia
from shared.config import ANTHROPIC_KEY, format_monto
from shared.db import get_empresa

log = logging.getLogger("prep.autofill.proposal")

TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "templates")

# Estable entre llamadas, asi que es lo que se cachea. Lo que cambia -- la
# licitacion y la empresa -- viaja en el mensaje.
SISTEMA = """Redactas propuestas tecnicas para contratacion publica peruana, \
bajo la Ley 32069 y su reglamento.

Como escribes:
- Concreto y verificable. El comite califica contra los terminos de referencia,
  no contra adjetivos: "cuadrilla de 4 operarios con supervisor a tiempo
  completo" puntua, "amplia experiencia y compromiso con la calidad" no.
- Solo con lo que la empresa acredita. No inventes certificaciones, obras,
  equipos ni personal: una propuesta que promete lo que no se puede sustentar
  se cae en la verificacion posterior y arrastra sancion del RNP.
- Si falta un dato, escribe la seccion con lo que hay y no rellenes el hueco
  con generalidades.
- Espanol de Peru, sin relleno corporativo."""

ESQUEMA = {
    "type": "object",
    "properties": {
        "secciones": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "titulo": {"type": "string"},
                    "contenido": {"type": "array", "items": {"type": "string"}},
                    "lista": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["titulo", "contenido", "lista"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["secciones"],
    "additionalProperties": False,
}


async def generar_propuesta_tecnica(propuesta_id: int, empresa_id: int,
                                     licitacion: dict, datos: dict) -> str | None:
    """Genera propuesta técnica completa usando Claude API.

    Secciones estándar:
    1. Resumen ejecutivo
    2. Objetivos
    3. Alcance del servicio
    4. Metodología
    5. Plan de trabajo
    6. Equipo técnico propuesto
    7. Equipamiento y recursos
    8. Experiencia similar
    9. Cronograma
    10. Valor agregado
    """
    empresa = await get_empresa(empresa_id)
    if not empresa:
        return None

    # Generar contenido con IA o plantilla
    if ANTHROPIC_KEY:
        contenido = await _generar_con_ia(licitacion, empresa, datos)
    else:
        contenido = _generar_plantilla(licitacion, empresa, datos)

    # Crear documento DOCX
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Portada
    doc.add_paragraph()
    doc.add_paragraph()
    titulo = doc.add_heading("PROPUESTA TECNICA", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run(f"\n{licitacion.get('nomenclatura', licitacion.get('id', ''))}").font.size = Pt(14)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"\n{licitacion['objeto'][:200]}").font.size = Pt(12)

    doc.add_paragraph()
    empresa_p = doc.add_paragraph()
    empresa_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    empresa_p.add_run(f"\nPresentada por:\n{empresa['razon_social']}\nRUC: {empresa['ruc']}").font.size = Pt(12)
    doc.add_page_break()

    # Secciones
    for seccion in contenido.get("secciones", []):
        doc.add_heading(seccion["titulo"], level=1)
        for parrafo in seccion.get("contenido", []):
            doc.add_paragraph(parrafo)
        if seccion.get("lista"):
            for item in seccion["lista"]:
                doc.add_paragraph(f"  - {item}")
        doc.add_paragraph()

    # Equipo técnico (tabla)
    if datos.get("equipo"):
        doc.add_heading("EQUIPO TECNICO PROPUESTO", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["Nombre", "Cargo", "Titulo", "Experiencia"]):
            table.rows[0].cells[i].text = h
        for m in datos["equipo"][:8]:
            row = table.add_row()
            row.cells[0].text = m.get("nombre_completo", "")
            row.cells[1].text = m.get("especialidad", "")
            row.cells[2].text = m.get("titulo_profesional", "")
            row.cells[3].text = f"{m.get('anos_experiencia', 0)} años"

    output_dir = os.path.join(TEMPLATES_DIR, "output")
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"propuesta_tecnica_{propuesta_id}.docx")
    doc.save(path)
    log.info(f"Propuesta técnica generada: {path}")
    return path


async def _generar_con_ia(licitacion: dict, empresa, datos: dict) -> dict:
    """Genera contenido de propuesta técnica con Claude API."""
    equipo_text = ""
    if datos.get("equipo"):
        equipo_text = "\n".join(
            f"- {m['nombre_completo']}: {m.get('titulo_profesional', '')} ({m.get('anos_experiencia', 0)} anos)"
            for m in datos["equipo"][:5]
        )

    exp_text = ""
    if datos.get("experiencia"):
        exp_text = "\n".join(
            f"- {e['objeto_contrato'][:80]} ({e.get('entidad_contratante', '')}): S/{e.get('monto') or 0:,.0f}"
            for e in datos["experiencia"][:5]
        )

    prompt = f"""Genera una propuesta técnica profesional para esta licitación pública peruana.

LICITACIÓN:
- Objeto: {licitacion['objeto']}
- Entidad: {licitacion['entidad']}
- Monto ref: {format_monto(licitacion.get('monto_referencial', 0)) if licitacion.get('monto_referencial') else 'No especificado'}

EMPRESA:
- {empresa['razon_social']} (RUC {empresa['ruc']})
- Rubros: {', '.join(empresa.get('rubros') or [])}

EQUIPO DISPONIBLE:
{equipo_text or 'No registrado'}

EXPERIENCIA:
{exp_text or 'No registrada'}

Escribe siete secciones, en este orden y con estos titulos exactos:
1. RESUMEN EJECUTIVO, 2. OBJETIVOS, 3. ALCANCE DEL SERVICIO, 4. METODOLOGIA,
5. PLAN DE TRABAJO, 6. EXPERIENCIA SIMILAR, 7. VALOR AGREGADO.

Maximo tres o cuatro parrafos por seccion."""

    try:
        contenido, _uso = await ia.pedir_json(SISTEMA, prompt, ESQUEMA)
        return contenido
    except Exception:
        # Se cae a la plantilla: no lleva IA, pero produce un documento
        # presentable. Quedarse sin propuesta tecnica es peor -- en un concurso
        # publico o una licitacion publica es el documento que puntua.
        log.exception("La propuesta tecnica con IA fallo")

    return _generar_plantilla(licitacion, empresa, datos)


def _generar_plantilla(licitacion: dict, empresa, datos: dict) -> dict:
    """Plantilla básica sin IA."""
    return {
        "secciones": [
            {
                "titulo": "1. RESUMEN EJECUTIVO",
                "contenido": [
                    (f"{empresa['razon_social']} presenta esta propuesta técnica para la contratación de: "
                    f"{licitacion['objeto']}."),
                    ("Nuestra empresa cuenta con la experiencia, el equipo técnico y los recursos "
                    "necesarios para ejecutar el presente servicio dentro de los plazos y condiciones "
                    "establecidos en las bases del procedimiento de selección."),
                ],
            },
            {
                "titulo": "2. OBJETIVOS",
                "contenido": ["Los objetivos de la presente propuesta son:"],
                "lista": [
                    "Cumplir con todos los requerimientos técnicos establecidos en las bases.",
                    "Entregar el servicio/bien dentro del plazo establecido.",
                    "Garantizar la calidad conforme a los estándares solicitados.",
                    "Proporcionar soporte y garantía post-entrega.",
                ],
            },
            {
                "titulo": "3. ALCANCE DEL SERVICIO",
                "contenido": [
                    (f"El alcance comprende la totalidad de lo solicitado en el objeto de contratación: "
                    f"{licitacion['objeto'][:200]}."),
                ],
            },
            {
                "titulo": "4. METODOLOGIA",
                "contenido": [
                    ("Nuestra metodología se basa en las mejores prácticas del sector, "
                    "con un enfoque estructurado en fases: planificación, ejecución, "
                    "control de calidad y entrega."),
                ],
            },
            {
                "titulo": "5. PLAN DE TRABAJO",
                "contenido": ["El plan de trabajo se estructura en las siguientes fases:"],
                "lista": [
                    "Fase 1: Planificación y coordinación inicial",
                    "Fase 2: Ejecución del servicio/entrega del bien",
                    "Fase 3: Control de calidad y ajustes",
                    "Fase 4: Entrega final y conformidad",
                ],
            },
        ],
    }
