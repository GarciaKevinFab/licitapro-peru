"""Motor de auto-fill — Llena anexos usando knowledge_base + datos empresa."""
import os
import logging
from datetime import date
from docx import Document
from shared.db import kb_get, kb_set, get_empresa, connection

log = logging.getLogger("prep.autofill")

TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "/app/templates")


# Campos requeridos por tipo de anexo
ANEXO_FIELDS = {
    "carta_presentacion": [
        ("legal", "representante_legal", "Nombre del representante legal"),
        ("legal", "dni_representante", "DNI del representante"),
        ("legal", "cargo_representante", "Cargo del representante"),
        ("legal", "razon_social", None),  # None = no preguntar, viene de empresa
        ("legal", "ruc", None),
        ("legal", "direccion_legal", "Dirección legal completa"),
        ("legal", "telefono", None),
        ("legal", "email", None),
    ],
    "declaracion_jurada": [
        ("legal", "representante_legal", None),
        ("legal", "dni_representante", None),
        ("legal", "razon_social", None),
        ("legal", "ruc", None),
        ("legal", "partida_registral", "Número de partida registral (SUNARP)"),
    ],
    "propuesta_economica": [
        ("financiero", "entidad_bancaria", "¿Banco donde tiene cuenta la empresa?"),
        ("financiero", "cuenta_corriente", "Número de cuenta corriente"),
        ("financiero", "cci", "Código de cuenta interbancario (CCI)"),
        ("financiero", "margen_defecto", "¿Margen mínimo aceptable sobre costo? (ej: 20%)"),
    ],
    "experiencia_postor": [
        # Experiencia se llena desde la tabla experiencia, no KB
    ],
    "equipo_tecnico": [
        # Equipo se llena desde tabla equipo_tecnico
    ],
}


async def autofill_propuesta(propuesta_id: int, empresa_id: int) -> dict:
    """
    Intenta llenar todos los campos de todos los anexos.
    Retorna: {
        "completados": int,
        "totales": int,
        "faltantes": [{"categoria", "clave", "pregunta"}],
        "datos": {campo: valor}
    }
    """
    empresa = await get_empresa(empresa_id)
    if not empresa:
        return {"completados": 0, "totales": 0, "faltantes": [], "datos": {}}

    datos = {}
    faltantes = []
    completados = 0
    totales = 0

    # Pre-cargar datos de la empresa en KB si no están
    empresa_fields = {
        "razon_social": empresa["razon_social"],
        "ruc": empresa["ruc"],
        "telefono": empresa["telefono"],
        "email": empresa["email"],
        "direccion_legal": empresa["direccion"],
        "representante_legal": empresa["representante_legal"],
        "dni_representante": empresa["dni_representante"],
        "cargo_representante": empresa.get("cargo_representante", "Gerente General"),
    }

    for clave, valor in empresa_fields.items():
        if valor:
            await kb_set(empresa_id, "legal", clave, str(valor), "tabla_empresas")

    # Recorrer todos los anexos y sus campos
    for anexo_tipo, fields in ANEXO_FIELDS.items():
        for cat, clave, pregunta_text in fields:
            totales += 1
            
            # Buscar en KB
            valor = await kb_get(empresa_id, cat, clave)
            if valor:
                datos[f"{anexo_tipo}.{clave}"] = valor
                completados += 1
                continue

            # Si no hay pregunta definida, significa que debería venir de empresa
            if pregunta_text is None:
                # Ya intentamos arriba, no lo tiene
                completados += 1  # Lo contamos como OK si viene de empresa
                datos[f"{anexo_tipo}.{clave}"] = empresa_fields.get(clave, "")
                continue

            # Faltante → necesita pregunta al usuario
            faltantes.append({
                "categoria": cat,
                "clave": clave,
                "pregunta": pregunta_text,
                "anexo": anexo_tipo,
            })

    # Cargar experiencia
    async with connection() as conn:
        experiencias = await conn.fetch(
            "SELECT * FROM experiencia WHERE empresa_id=$1 ORDER BY fecha_fin DESC",
            empresa_id,
        )
        datos["experiencia_count"] = len(experiencias)
        datos["experiencias"] = [dict(e) for e in experiencias]

        # Cargar equipo técnico
        equipo = await conn.fetch(
            "SELECT * FROM equipo_tecnico WHERE empresa_id=$1 AND disponible=TRUE",
            empresa_id,
        )
        datos["equipo_count"] = len(equipo)
        datos["equipo"] = [dict(e) for e in equipo]

    return {
        "completados": completados,
        "totales": totales,
        "faltantes": faltantes,
        "datos": datos,
    }


async def generar_carta_presentacion(datos: dict, licitacion: dict) -> str:
    """Genera carta de presentación usando template o desde cero."""
    
    doc = Document()
    
    # Encabezado
    doc.add_paragraph(f"Lima, {date.today().strftime('%d de %B de %Y')}")
    doc.add_paragraph("")
    doc.add_paragraph("Señores")
    doc.add_paragraph(f"{licitacion.get('entidad', '')}")
    doc.add_paragraph("Presente.-")
    doc.add_paragraph("")
    
    # Referencia
    ref = licitacion.get("nomenclatura", licitacion.get("id", ""))
    doc.add_paragraph(f"Ref.: {ref}")
    doc.add_paragraph(f"Objeto: {licitacion.get('objeto', '')}")
    doc.add_paragraph("")
    
    # Cuerpo
    razon = datos.get("carta_presentacion.razon_social", "")
    ruc = datos.get("carta_presentacion.ruc", "")
    rep = datos.get("carta_presentacion.representante_legal", "")
    dni = datos.get("carta_presentacion.dni_representante", "")
    
    doc.add_paragraph(
        f"De nuestra consideración:\n\n"
        f"El que suscribe, {rep}, identificado con DNI N° {dni}, "
        f"representante legal de {razon}, con RUC N° {ruc}, "
        f"en atención al procedimiento de selección de la referencia, "
        f"presenta la siguiente propuesta técnica y económica para la "
        f"contratación del servicio/bien indicado.\n\n"
        f"Declaramos que nuestra propuesta cumple con los requisitos "
        f"establecidos en las bases del procedimiento y nos comprometemos "
        f"a ejecutar el contrato en los términos y condiciones establecidos."
    )
    
    doc.add_paragraph("")
    doc.add_paragraph("Atentamente,")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("______________________________")
    doc.add_paragraph(f"{rep}")
    doc.add_paragraph(f"DNI: {dni}")
    doc.add_paragraph(f"{razon}")
    doc.add_paragraph(f"RUC: {ruc}")
    
    # Guardar
    output_dir = os.getenv("EXPEDIENTES_DIR", "/app/data/expedientes")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"carta_presentacion_{licitacion['id']}.docx")
    doc.save(output_path)
    
    return output_path


async def generar_declaracion_jurada(datos: dict, licitacion: dict) -> str:
    """Genera declaración jurada estándar."""
    doc = Document()
    
    razon = datos.get("declaracion_jurada.razon_social", "")
    ruc = datos.get("declaracion_jurada.ruc", "")
    rep = datos.get("declaracion_jurada.representante_legal", "")
    dni = datos.get("declaracion_jurada.dni_representante", "")
    
    doc.add_heading("DECLARACIÓN JURADA", level=1)
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Yo, {rep}, identificado(a) con DNI N° {dni}, "
        f"representante legal de {razon}, con RUC N° {ruc}, "
        f"DECLARO BAJO JURAMENTO que:"
    )
    doc.add_paragraph("")
    
    declaraciones = [
        "No tengo impedimento para participar en el procedimiento de selección ni para contratar con el Estado.",
        "Conozco, acepto y me someto a las bases, condiciones y reglas del procedimiento de selección.",
        "Soy responsable de la veracidad de los documentos e información que presento.",
        "Me comprometo a mantener la oferta durante el procedimiento de selección y a suscribir el contrato en caso de obtener la buena pro.",
        "Conozco las sanciones contenidas en la Ley de Contrataciones del Estado y su Reglamento.",
        "Participamos en el presente procedimiento de selección en forma independiente, sin mediar consulta, comunicación, acuerdo, arreglo o convenio con ningún proveedor.",
    ]
    
    for i, decl in enumerate(declaraciones, 1):
        doc.add_paragraph(f"{i}. {decl}")
    
    doc.add_paragraph("")
    doc.add_paragraph(f"Lima, {date.today().strftime('%d de %B de %Y')}")
    doc.add_paragraph("")
    doc.add_paragraph("______________________________")
    doc.add_paragraph(f"{rep}")
    doc.add_paragraph(f"DNI: {dni}")
    
    output_dir = os.getenv("EXPEDIENTES_DIR", "/app/data/expedientes")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"declaracion_jurada_{licitacion['id']}.docx")
    doc.save(output_path)
    
    return output_path
