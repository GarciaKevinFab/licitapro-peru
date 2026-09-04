"""Anexos estandar del expediente.

QUE APORTA ESTE MODULO Y QUE NO

  De los catorce anexos de la lista, `zip_builder` ya genera cinco por su
  cuenta (carta, declaracion jurada, experiencia, propuesta tecnica y
  economica) y cuatro no los puede generar nadie porque son documentos que
  emite un tercero: constancia del RNP, vigencia de poder, constancia de no
  estar inhabilitado y promesa de consorcio.

  Lo que solo esta aqui son TRES: la declaracion jurada de plazo de entrega, la
  carta de compromiso del personal clave y el pacto de integridad. Los tres se
  exigen de forma habitual y los tres se pueden escribir con lo que ya hay en
  la base. Por eso `generar_anexos_complementarios` genera esos tres y no llama
  a `llenar_anexos`: llamarla duplicaria los cinco que el ZIP ya trae.

  `llenar_anexos` se conserva para el dia que haga falta el inventario completo
  con su estado; hoy el expediente se arma con la funcion de abajo.
"""
import logging
import os

from prep_bot.document_gen import (
    generar_carta_presentacion,
    generar_declaracion_jurada,
    generar_experiencia_postor,
)
from shared.db import connection, get_empresa

log = logging.getLogger("prep.autofill.annexes")

# Anexos estándar en licitaciones peruanas (Ley 30225)
ANEXOS_ESTANDAR = [
    {"num": 1, "nombre": "Carta de Presentación", "tipo": "docx", "auto": True},
    {"num": 2, "nombre": "Declaración Jurada", "tipo": "docx", "auto": True},
    {"num": 3, "nombre": "Declaración Jurada de Plazo de Entrega", "tipo": "docx", "auto": True},
    {"num": 4, "nombre": "Experiencia del Postor", "tipo": "docx", "auto": True},
    {"num": 5, "nombre": "Carta de Compromiso del Personal Clave", "tipo": "docx", "auto": True},
    {"num": 6, "nombre": "Promesa de Consorcio", "tipo": "docx", "auto": False},  # Solo si aplica
    {"num": 7, "nombre": "Propuesta Económica", "tipo": "xlsx", "auto": True},
    {"num": 8, "nombre": "Constancia RNP + Vigencia de Poder", "tipo": "pdf", "auto": False},
    {"num": 9, "nombre": "Constancia de No Estar Inhabilitado", "tipo": "pdf", "auto": False},
    {"num": 10, "nombre": "Pacto de Integridad", "tipo": "docx", "auto": True},
    {"num": 11, "nombre": "Propuesta Técnica", "tipo": "docx", "auto": True},
    {"num": 12, "nombre": "Equipo Técnico (CVs)", "tipo": "docx", "auto": True},
    {"num": 13, "nombre": "Plan de Trabajo", "tipo": "docx", "auto": True},
    {"num": 14, "nombre": "Equipamiento Disponible", "tipo": "docx", "auto": True},
]


async def generar_anexos_complementarios(
        propuesta_id: int, empresa_id: int, licitacion: dict,
        datos: dict) -> list[tuple[str, str]]:
    """Los tres anexos que el resto del expediente no cubre.

    Devuelve pares (nombre dentro del ZIP, ruta en disco). Cada uno se intenta
    por separado: que falle el pacto de integridad no puede dejar fuera la
    declaracion jurada de plazo, porque son documentos independientes y el
    expediente vale mas con dos que con ninguno.
    """
    empresa = await get_empresa(empresa_id)
    if not empresa:
        return []

    salida = []
    output_dir = os.path.join(os.getenv("TEMPLATES_DIR", "templates"), "output")
    os.makedirs(output_dir, exist_ok=True)

    async with connection() as conn:
        tiene_equipo = await conn.fetchval(
            """SELECT EXISTS(SELECT 1 FROM equipo_tecnico
                              WHERE empresa_id=$1 AND disponible=TRUE)""",
            empresa_id)

    tareas = [
        ("06_DJ_Plazo_Entrega.docx",
         _generar_dj_plazo(propuesta_id, empresa, licitacion, datos, output_dir)),
        ("08_Pacto_Integridad.docx",
         _generar_pacto_integridad(propuesta_id, empresa, licitacion, datos,
                                   output_dir)),
    ]
    # Sin profesionales registrados, la carta de compromiso saldria sin nadie
    # que se comprometa. Un anexo vacio en el expediente se lee peor que un
    # anexo ausente: parece que se respondio y no se respondio.
    if tiene_equipo:
        tareas.insert(1, (
            "07_Compromiso_Personal_Clave.docx",
            _generar_compromiso_personal(propuesta_id, empresa_id, licitacion,
                                         output_dir)))

    for nombre, tarea in tareas:
        try:
            ruta = await tarea
            if ruta:
                salida.append((nombre, ruta))
        except Exception as e:
            log.error("El anexo %s fallo: %s", nombre, e, exc_info=True)

    return salida


async def llenar_anexos(propuesta_id: int, empresa_id: int,
                         licitacion: dict, datos: dict) -> dict:
    """Llena todos los anexos posibles automáticamente.

    Returns: {completados, total, detalle: [{anexo, estado, path}]}
    """
    empresa = await get_empresa(empresa_id)
    completados = 0
    detalle = []
    templates_dir = os.getenv("TEMPLATES_DIR", "templates")
    output_dir = os.path.join(templates_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    for anexo in ANEXOS_ESTANDAR:
        try:
            if not anexo["auto"]:
                detalle.append({
                    "anexo": f"Anexo {anexo['num']}: {anexo['nombre']}",
                    "estado": "manual",
                    "nota": "Requiere documento externo (PDF de entidad)",
                })
                continue

            path = await _generar_anexo(
                anexo, propuesta_id, empresa_id, licitacion, datos, empresa, output_dir
            )
            if path:
                completados += 1
                detalle.append({
                    "anexo": f"Anexo {anexo['num']}: {anexo['nombre']}",
                    "estado": "completado",
                    "path": path,
                })
            else:
                detalle.append({
                    "anexo": f"Anexo {anexo['num']}: {anexo['nombre']}",
                    "estado": "pendiente",
                    "nota": "Faltan datos",
                })

        except Exception as e:
            log.error(f"Error generando Anexo {anexo['num']}: {e}")
            detalle.append({
                "anexo": f"Anexo {anexo['num']}: {anexo['nombre']}",
                "estado": "error",
                "error": str(e),
            })

    return {
        "completados": completados,
        "total": len(ANEXOS_ESTANDAR),
        "detalle": detalle,
    }


async def _generar_anexo(anexo: dict, propuesta_id: int, empresa_id: int,
                          licitacion: dict, datos: dict, empresa, output_dir: str) -> str | None:
    """Genera un anexo específico."""

    if anexo["num"] == 1:
        return await generar_carta_presentacion(propuesta_id, empresa_id, licitacion)

    elif anexo["num"] == 2:
        return await generar_declaracion_jurada(propuesta_id, empresa_id, licitacion)

    elif anexo["num"] == 3:
        return await _generar_dj_plazo(propuesta_id, empresa, licitacion, datos, output_dir)

    elif anexo["num"] == 4:
        return await generar_experiencia_postor(propuesta_id, empresa_id, licitacion)

    elif anexo["num"] == 5:
        return await _generar_compromiso_personal(propuesta_id, empresa_id, licitacion, output_dir)

    elif anexo["num"] == 10:
        return await _generar_pacto_integridad(propuesta_id, empresa, licitacion, datos, output_dir)

    elif anexo["num"] in (11, 12, 13, 14):
        # Estos se generan en otros módulos (proposal.py, etc.)
        path = os.path.join(output_dir, f"anexo_{anexo['num']}_{propuesta_id}.docx")
        if os.path.exists(path):
            return path
        return None

    return None


async def _generar_dj_plazo(propuesta_id, empresa, licitacion, datos, output_dir) -> str:
    """Genera Declaración Jurada de plazo de entrega."""
    from datetime import date

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("DECLARACION JURADA DE PLAZO DE ENTREGA", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    representante = datos.get("legal.representante_legal", empresa.get("representante_legal", "[REPRESENTANTE]"))

    doc.add_paragraph(
        f"Yo, {representante}, representante legal de {empresa['razon_social']}, "
        f"con RUC N° {empresa['ruc']}, declaro bajo juramento que nos comprometemos a "
        f"entregar el bien/servicio objeto de la contratación dentro del plazo establecido "
        f"en las Bases del proceso {licitacion.get('nomenclatura', licitacion.get('id', ''))}."
    )

    doc.add_paragraph()
    doc.add_paragraph(f"{date.today().strftime('%d de %B de %Y')}")
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"{representante}")
    doc.add_paragraph(f"{empresa['razon_social']}")

    path = os.path.join(output_dir, f"dj_plazo_{propuesta_id}.docx")
    doc.save(path)
    return path


async def _generar_compromiso_personal(propuesta_id, empresa_id, licitacion, output_dir) -> str:
    """Genera Carta de Compromiso del Personal Clave."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    async with connection() as conn:
        equipo = await conn.fetch(
            "SELECT * FROM equipo_tecnico WHERE empresa_id=$1 AND disponible=TRUE LIMIT 5",
            empresa_id,
        )

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("CARTA DE COMPROMISO DEL PERSONAL CLAVE", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if equipo:
        for m in equipo:
            doc.add_paragraph(
                f"Yo, {m['nombre_completo']}, identificado(a) con DNI N° {m['dni'] or '[DNI]'}, "
                f"profesional en {m['titulo_profesional'] or '[TITULO]'}, "
                f"me comprometo a participar como parte del equipo técnico "
                f"para el proceso de selección indicado."
            )
            doc.add_paragraph()
    else:
        doc.add_paragraph("[Pendiente: agregar profesionales al equipo técnico]")

    path = os.path.join(output_dir, f"compromiso_personal_{propuesta_id}.docx")
    doc.save(path)
    return path


async def _generar_pacto_integridad(propuesta_id, empresa, licitacion, datos, output_dir) -> str:
    """Genera Pacto de Integridad."""
    from datetime import date

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("PACTO DE INTEGRIDAD", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    representante = datos.get("legal.representante_legal", empresa.get("representante_legal", "[REPRESENTANTE]"))

    doc.add_paragraph(
        f"Yo, {representante}, representante legal de {empresa['razon_social']}, "
        f"con RUC N° {empresa['ruc']}, manifiesto mi compromiso de conducirme con "
        f"integridad, veracidad, transparencia y legalidad durante el presente proceso "
        f"de contratación y, de ser el caso, durante la ejecución contractual."
    )
    doc.add_paragraph()
    doc.add_paragraph("En tal sentido, me comprometo a:")
    compromisos = [
        "No ofrecer ni conceder beneficios indebidos a funcionario público alguno.",
        "No celebrar acuerdos con otros postores para manipular los resultados.",
        "Cumplir con la legislación vigente en materia de contrataciones del Estado.",
        "Denunciar cualquier acto de corrupción del que tome conocimiento.",
    ]
    for c in compromisos:
        doc.add_paragraph(f"  - {c}")

    doc.add_paragraph()
    doc.add_paragraph(f"{date.today().strftime('%d de %B de %Y')}")
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"{representante}")
    doc.add_paragraph(f"{empresa['razon_social']}")

    path = os.path.join(output_dir, f"pacto_integridad_{propuesta_id}.docx")
    doc.save(path)
    return path
