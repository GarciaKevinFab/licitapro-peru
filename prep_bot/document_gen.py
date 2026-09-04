"""Document Generator — Genera documentos DOCX/PDF para propuestas."""
import logging
import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from shared.config import format_fecha, format_monto
from shared.db import connection, get_empresa
from shared.firma_manager import insertar_imagen_en_docx, obtener_firma

log = logging.getLogger("prep.docgen")

TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "templates")


def _set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"


async def generar_carta_presentacion(propuesta_id: int, empresa_id: int, licitacion: dict) -> str:
    """Genera Anexo 1: Carta de Presentación del postor."""
    empresa = await get_empresa(empresa_id)
    datos = await _recopilar_datos(empresa_id)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Encabezado
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{datos.get('legal.domicilio_legal', empresa.get('direccion', ''))}\n")
    p.add_run(f"{date.today().strftime('%d de %B de %Y')}")

    doc.add_paragraph()  # Espacio

    # Destinatario
    doc.add_paragraph("Señores")
    doc.add_paragraph(f"{licitacion['entidad']}").bold = True
    doc.add_paragraph("Presente.-")
    doc.add_paragraph()

    # Referencia
    ref = doc.add_paragraph()
    ref.add_run("Ref: ").bold = True
    ref.add_run(f"{licitacion.get('nomenclatura', licitacion['id'])}")
    doc.add_paragraph()

    # Cuerpo
    representante = datos.get("legal.representante_legal", empresa.get("representante_legal", "[REPRESENTANTE]"))
    cargo = datos.get("legal.cargo_representante", "Gerente General")
    razon = empresa["razon_social"]
    ruc = empresa["ruc"]

    body = (
        f"Estimados señores:\n\n"
        f"El que suscribe, {representante}, identificado con DNI N° "
        f"{datos.get('legal.dni_representante', empresa.get('dni_representante', '[DNI]'))}, "
        f"en calidad de {cargo} y representante legal de {razon}, "
        f"identificada con RUC N° {ruc}, domiciliada en "
        f"{datos.get('legal.domicilio_legal', empresa.get('direccion', '[DIRECCION]'))}, "
        f"presenta a ustedes la propuesta técnica y económica para la contratación de:\n\n"
        f"\"{licitacion['objeto']}\"\n\n"
        f"Para lo cual declaramos que cumplimos con los requisitos de calificación y "
        f"demás condiciones establecidas en las Bases del presente procedimiento de selección."
    )
    doc.add_paragraph(body)

    doc.add_paragraph()
    doc.add_paragraph("Atentamente,")
    doc.add_paragraph()

    # Insertar firma del representante (imagen)
    firma_path = await obtener_firma(empresa_id, 'firma')
    if firma_path:
        insertar_imagen_en_docx(doc, firma_path, width_cm=4.0)
    else:
        doc.add_paragraph()
        p_line = doc.add_paragraph()
        p_line.add_run("_" * 40)

    doc.add_paragraph(f"{representante}")
    doc.add_paragraph(f"{cargo}")
    doc.add_paragraph(f"{razon}")
    doc.add_paragraph(f"RUC: {ruc}")

    # Insertar sello si existe
    sello_path = await obtener_firma(empresa_id, 'sello')
    if sello_path:
        insertar_imagen_en_docx(doc, sello_path, width_cm=3.0)

    path = os.path.join(TEMPLATES_DIR, "output", f"carta_presentacion_{propuesta_id}.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    log.info(f"Carta de presentación generada: {path}")
    return path


async def generar_declaracion_jurada(propuesta_id: int, empresa_id: int, licitacion: dict) -> str:
    """Genera Anexo 2: Declaración Jurada del postor."""
    empresa = await get_empresa(empresa_id)
    datos = await _recopilar_datos(empresa_id)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("DECLARACION JURADA", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref.add_run(f"({licitacion.get('nomenclatura', licitacion['id'])})")
    doc.add_paragraph()

    representante = datos.get("legal.representante_legal", empresa.get("representante_legal", "[REPRESENTANTE]"))
    razon = empresa["razon_social"]
    ruc = empresa["ruc"]

    declaraciones = [
        f"Que {razon}, con RUC N° {ruc}, no se encuentra incursa en ninguno de los impedimentos para contratar con el Estado previstos en el artículo 11 de la Ley de Contrataciones del Estado.",
        "Que conocemos, aceptamos y nos sometemos a las condiciones y procedimientos del presente proceso de selección.",
        "Que somos responsables de la veracidad de los documentos e información que presentamos.",
        "Que nos comprometemos a mantener nuestra oferta durante el proceso de selección y a suscribir el contrato en caso de resultar favorecidos con la Buena Pro.",
        "Que no tenemos impedimento para participar en el proceso de selección ni para contratar con el Estado.",
        "Que conocemos las sanciones contenidas en la Ley de Contrataciones del Estado y su Reglamento.",
    ]

    for i, decl in enumerate(declaraciones, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(decl)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()
    doc.add_paragraph(f"{datos.get('legal.domicilio_legal', '')}, {date.today().strftime('%d de %B de %Y')}")
    doc.add_paragraph()

    # Firma del representante
    firma_path = await obtener_firma(empresa_id, 'firma')
    if firma_path:
        insertar_imagen_en_docx(doc, firma_path, width_cm=4.0)
    else:
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)

    doc.add_paragraph(f"{representante}")
    doc.add_paragraph("Representante Legal")
    doc.add_paragraph(f"{razon}")

    sello_path = await obtener_firma(empresa_id, 'sello')
    if sello_path:
        insertar_imagen_en_docx(doc, sello_path, width_cm=3.0)

    path = os.path.join(TEMPLATES_DIR, "output", f"declaracion_jurada_{propuesta_id}.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


async def generar_experiencia_postor(propuesta_id: int, empresa_id: int, licitacion: dict) -> str:
    """Genera Anexo 4: Experiencia del Postor."""
    empresa = await get_empresa(empresa_id)

    async with connection() as conn:
        experiencias = await conn.fetch(
            "SELECT * FROM experiencia WHERE empresa_id=$1 ORDER BY monto DESC LIMIT 10",
            empresa_id,
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    doc.add_heading("EXPERIENCIA DEL POSTOR", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Empresa: {empresa['razon_social']} | RUC: {empresa['ruc']}")
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["N°", "Entidad Contratante", "Objeto del Contrato", "Monto (S/)", "Fecha Inicio", "Fecha Fin"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=9)

    for idx, exp in enumerate(experiencias, 1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(idx), size=9)
        _set_cell_text(row.cells[1], exp["entidad_contratante"] or "", size=9)
        _set_cell_text(row.cells[2], exp["objeto_contrato"] or "", size=9)
        _set_cell_text(row.cells[3], format_monto(exp["monto"]) if exp["monto"] else "—", size=9)
        _set_cell_text(row.cells[4], format_fecha(exp["fecha_inicio"]), size=9)
        _set_cell_text(row.cells[5], format_fecha(exp["fecha_fin"]), size=9)

    if experiencias:
        total = sum(e["monto"] for e in experiencias if e["monto"])
        doc.add_paragraph(f"\nMonto acumulado: {format_monto(total)}")

    path = os.path.join(TEMPLATES_DIR, "output", f"experiencia_postor_{propuesta_id}.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


async def generar_propuesta_economica(propuesta_id: int, empresa_id: int,
                                       licitacion: dict, precio: float) -> str:
    """Genera propuesta económica."""
    empresa = await get_empresa(empresa_id)
    datos = await _recopilar_datos(empresa_id)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("PROPUESTA ECONOMICA", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Ref: {licitacion.get('nomenclatura', licitacion['id'])}")
    doc.add_paragraph()

    doc.add_paragraph(f"Señores: {licitacion['entidad']}")
    doc.add_paragraph()

    representante = datos.get("legal.representante_legal", empresa.get("representante_legal", ""))
    body = (
        f"Es grato dirigirnos a ustedes para presentar nuestra propuesta económica "
        f"por el monto total de {format_monto(precio)} (incluido IGV) para la contratación de:\n\n"
        f"\"{licitacion['objeto']}\"\n\n"
        f"Esta oferta incluye todos los tributos, seguros, transportes, inspecciones, "
        f"pruebas, y de ser el caso, los costos laborales conforme a la legislación vigente, "
        f"así como cualquier otro concepto que pueda tener incidencia sobre el costo."
    )
    doc.add_paragraph(body)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Concepto", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Monto", bold=True)
    _set_cell_text(table.rows[1].cells[0], "Subtotal sin IGV")
    _set_cell_text(table.rows[1].cells[1], format_monto(precio / 1.18))
    _set_cell_text(table.rows[2].cells[0], "IGV (18%)")
    _set_cell_text(table.rows[2].cells[1], format_monto(precio - precio / 1.18))
    _set_cell_text(table.rows[3].cells[0], "TOTAL", bold=True)
    _set_cell_text(table.rows[3].cells[1], format_monto(precio), bold=True)

    doc.add_paragraph()
    doc.add_paragraph(f"{date.today().strftime('%d de %B de %Y')}")
    doc.add_paragraph()

    firma_path = await obtener_firma(empresa_id, 'firma')
    if firma_path:
        insertar_imagen_en_docx(doc, firma_path, width_cm=4.0)
    else:
        firma = doc.add_paragraph()
        firma.add_run("_" * 40)

    doc.add_paragraph(f"{representante}")
    doc.add_paragraph(f"{empresa['razon_social']} | RUC: {empresa['ruc']}")

    sello_path = await obtener_firma(empresa_id, 'sello')
    if sello_path:
        insertar_imagen_en_docx(doc, sello_path, width_cm=3.0)

    path = os.path.join(TEMPLATES_DIR, "output", f"propuesta_economica_{propuesta_id}.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


async def _recopilar_datos(empresa_id: int) -> dict:
    """Recopila todos los datos de KB para una empresa."""
    async with connection() as conn:
        rows = await conn.fetch(
            "SELECT categoria, clave, valor FROM knowledge_base WHERE empresa_id=$1",
            empresa_id,
        )
    return {f"{r['categoria']}.{r['clave']}": r["valor"] for r in rows}
