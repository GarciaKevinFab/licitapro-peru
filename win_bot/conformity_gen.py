"""Conformity Generator — Genera actas de conformidad del servicio."""
import os
import logging
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from shared.db import connection, get_empresa
from shared.config import format_monto, format_fecha

log = logging.getLogger("win.conformity")

TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "templates")


async def generar_acta_conformidad(contrato_id: int, observaciones: str = "") -> str:
    """Genera acta de conformidad del servicio/bien entregado.

    Returns: path al archivo generado
    """
    async with connection() as conn:
        contrato = await conn.fetchrow(
            """SELECT c.*, l.entidad, l.objeto, l.nomenclatura,
                      e.razon_social, e.ruc, e.representante_legal
            FROM contratos c
            JOIN licitaciones l ON c.licitacion_id = l.id
            JOIN empresas e ON c.empresa_id = e.id
            WHERE c.id=$1""",
            contrato_id,
        )
        if not contrato:
            return None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Título
    titulo = doc.add_heading("ACTA DE CONFORMIDAD", level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.add_run(f"N° {contrato_id:04d}-{date.today().year}")

    doc.add_paragraph()

    # Datos del contrato
    doc.add_paragraph(f"En la ciudad de ____________, a los {date.today().day} días del mes "
                      f"de {date.today().strftime('%B')} de {date.today().year}, se procede "
                      f"a emitir la presente Acta de Conformidad respecto a:")

    doc.add_paragraph()

    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    datos = [
        ("Procedimiento:", contrato.get("nomenclatura") or f"Contrato #{contrato_id}"),
        ("Entidad:", contrato["entidad"]),
        ("Contratista:", f"{contrato['razon_social']} (RUC: {contrato['ruc']})"),
        ("Objeto:", contrato["objeto"][:200]),
        ("N° Contrato:", contrato.get("numero_contrato") or "—"),
        ("Monto:", format_monto(contrato["monto_adjudicado"]) if contrato.get("monto_adjudicado") else "—"),
        ("Fecha conformidad:", format_fecha(date.today())),
    ]
    for i, (label, value) in enumerate(datos):
        table.rows[i].cells[0].text = label
        for p in table.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Declaración de conformidad
    doc.add_paragraph(
        "Por medio de la presente, se deja constancia que el contratista ha cumplido "
        "con la prestación del servicio/entrega del bien de acuerdo con los términos "
        "de referencia/especificaciones técnicas y las condiciones establecidas en el "
        "contrato, en concordancia con lo dispuesto en el artículo 143 del Reglamento "
        "de la Ley de Contrataciones del Estado."
    )

    doc.add_paragraph()

    # Observaciones
    if observaciones:
        doc.add_heading("Observaciones:", level=2)
        doc.add_paragraph(observaciones)
        doc.add_paragraph()

    # Conclusión
    doc.add_paragraph(
        "En señal de conformidad, se suscribe la presente acta."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Firmas
    firma_table = doc.add_table(rows=4, cols=2)
    firma_table.rows[0].cells[0].text = "_" * 30
    firma_table.rows[0].cells[1].text = "_" * 30
    firma_table.rows[1].cells[0].text = "Área Usuaria"
    firma_table.rows[1].cells[1].text = contrato["representante_legal"] or "Contratista"
    firma_table.rows[2].cells[0].text = contrato["entidad"]
    firma_table.rows[2].cells[1].text = contrato["razon_social"]
    firma_table.rows[3].cells[0].text = ""
    firma_table.rows[3].cells[1].text = f"RUC: {contrato['ruc']}"

    # Guardar
    output_dir = os.path.join(TEMPLATES_DIR, "output", "conformidades")
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"acta_conformidad_{contrato_id}.docx")
    doc.save(path)

    # Actualizar estado del contrato
    async with connection() as conn:
        await conn.execute(
            "UPDATE contratos SET estado='conformidad' WHERE id=$1 AND estado='entregado'",
            contrato_id,
        )
        # Marcar plazo de conformidad como completado
        await conn.execute(
            """UPDATE plazos SET completado=TRUE
            WHERE contrato_id=$1 AND tipo='conformidad' AND completado=FALSE""",
            contrato_id,
        )

    log.info(f"Acta de conformidad generada: {path}")
    return path


async def generar_informe_entrega(contrato_id: int, descripcion_entrega: str) -> str:
    """Genera informe de entrega/avance parcial."""
    async with connection() as conn:
        contrato = await conn.fetchrow(
            """SELECT c.*, l.entidad, l.objeto, l.nomenclatura,
                      e.razon_social, e.ruc
            FROM contratos c
            JOIN licitaciones l ON c.licitacion_id = l.id
            JOIN empresas e ON c.empresa_id = e.id
            WHERE c.id=$1""",
            contrato_id,
        )
        if not contrato:
            return None

    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    doc.add_heading("INFORME DE ENTREGA", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph(f"Fecha: {format_fecha(date.today())}")
    doc.add_paragraph(f"Contrato: {contrato.get('numero_contrato', '—')}")
    doc.add_paragraph(f"Entidad: {contrato['entidad']}")
    doc.add_paragraph(f"Objeto: {contrato['objeto'][:200]}")
    doc.add_paragraph()

    doc.add_heading("Descripción de la entrega:", level=2)
    doc.add_paragraph(descripcion_entrega)
    doc.add_paragraph()

    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"{contrato['razon_social']}")
    doc.add_paragraph(f"RUC: {contrato['ruc']}")

    output_dir = os.path.join(TEMPLATES_DIR, "output", "entregas")
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"informe_entrega_{contrato_id}_{date.today().isoformat()}.docx")
    doc.save(path)

    # Actualizar estado
    async with connection() as conn:
        await conn.execute(
            "UPDATE contratos SET estado='entregado' WHERE id=$1 AND estado IN ('en_ejecucion', 'contrato_firmado')",
            contrato_id,
        )

    log.info(f"Informe de entrega generado: {path}")
    return path
