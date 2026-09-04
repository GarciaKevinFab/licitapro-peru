"""Invoice Generator — Genera facturas para cobro."""
import logging
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from shared import fechas
from shared.config import format_fecha, format_monto
from shared.db import connection
from shared.firma_manager import insertar_imagen_en_docx, obtener_firma

log = logging.getLogger("win.invoice")

TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", "templates")


async def generar_factura(contrato_id: int, monto: float, concepto: str,
                           numero_factura: str | None = None) -> str:
    """Genera documento de factura para cobro.

    Returns: path al archivo generado
    """
    async with connection() as conn:
        contrato = await conn.fetchrow(
            """SELECT c.*, l.entidad, l.objeto, l.nomenclatura, e.razon_social, e.ruc,
                      e.representante_legal, e.direccion, e.email
            FROM contratos c
            JOIN licitaciones l ON c.licitacion_id = l.id
            JOIN empresas e ON c.empresa_id = e.id
            WHERE c.id=$1""",
            contrato_id,
        )
        if not contrato:
            return None

    # Generar número de factura si no se proporcionó
    if not numero_factura:
        numero_factura = f"F001-{contrato_id:05d}"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Encabezado empresa
    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = enc.add_run(f"{contrato['razon_social']}\n")
    run.bold = True
    run.font.size = Pt(16)
    enc.add_run(f"RUC: {contrato['ruc']}\n").font.size = Pt(11)
    enc.add_run(f"{contrato.get('direccion', '')}\n").font.size = Pt(10)

    # Título
    #
    # DECIA "FACTURA ELECTRONICA" Y NO LO ES.
    #
    # Una factura electrónica en Perú es un XML firmado que se emite ante SUNAT,
    # directamente o a través de un OSE/PSE, y que devuelve una CDR. Esto es un
    # DOCX. La facturación electrónica quedó fuera de alcance a propósito -- lo
    # dice `web/contratos.py` -- así que el documento no puede llamarse como
    # aquello que el producto declara no hacer.
    #
    # El riesgo no es cosmético: un proveedor que ve "FACTURA ELECTRONICA N° 12"
    # en un Word se lo manda a la entidad creyendo que ya facturó. La entidad no
    # lo acepta, el plazo legal de pago no empieza a correr, y él se entera
    # semanas después preguntando por qué no le pagan.
    doc.add_paragraph()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"PROFORMA DE COBRO\nN° {numero_factura}")
    run.bold = True
    run.font.size = Pt(14)

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aclaracion = nota.add_run(
        "Documento informativo. NO es un comprobante de pago electrónico y no "
        "sustituye a la factura que debe emitirse ante SUNAT.")
    aclaracion.italic = True
    aclaracion.font.size = Pt(9)

    doc.add_paragraph()

    # Datos del cliente
    table_info = doc.add_table(rows=5, cols=2)
    table_info.style = "Table Grid"
    datos = [
        ("Fecha de emisión:", format_fecha(fechas.hoy())),
        ("Señor(es):", contrato["entidad"]),
        ("Referencia:", contrato.get("nomenclatura") or f"Contrato #{contrato_id}"),
        ("N° Contrato:", contrato.get("numero_contrato") or "—"),
        ("Condición:", "Contado"),
    ]
    for i, (label, value) in enumerate(datos):
        table_info.rows[i].cells[0].text = label
        table_info.rows[i].cells[0].paragraphs[0].runs[0].bold = True if table_info.rows[i].cells[0].paragraphs[0].runs else False
        table_info.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Detalle
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Cantidad", "Descripción", "P. Unitario", "Total"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True

    row = table.add_row()
    row.cells[0].text = "1"
    row.cells[1].text = f"{concepto}\n({contrato['objeto'][:150]})"
    subtotal = monto / 1.18
    row.cells[2].text = format_monto(subtotal)
    row.cells[3].text = format_monto(subtotal)

    doc.add_paragraph()

    # Totales
    total_table = doc.add_table(rows=3, cols=2)
    total_table.style = "Table Grid"
    totales = [
        ("Sub Total:", format_monto(subtotal)),
        ("IGV (18%):", format_monto(monto - subtotal)),
        ("TOTAL:", format_monto(monto)),
    ]
    for i, (label, value) in enumerate(totales):
        total_table.rows[i].cells[0].text = label
        total_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Datos bancarios
    banco_section = doc.add_paragraph()
    banco_section.add_run("DATOS BANCARIOS PARA PAGO:\n").bold = True

    async with connection() as conn:
        from shared.db import kb_get
        banco = await kb_get(contrato["empresa_id"], "financiero", "entidad_bancaria")
        cuenta = await kb_get(contrato["empresa_id"], "financiero", "cuenta_corriente")
        cci = await kb_get(contrato["empresa_id"], "financiero", "cci")

    if banco:
        doc.add_paragraph(f"Banco: {banco}")
    if cuenta:
        doc.add_paragraph(f"Cuenta Corriente: {cuenta}")
    if cci:
        doc.add_paragraph(f"CCI: {cci}")

    doc.add_paragraph(f"\nTitular: {contrato['razon_social']}")
    doc.add_paragraph(f"RUC: {contrato['ruc']}")

    # Firma y sello
    doc.add_paragraph()
    firma_path = await obtener_firma(contrato["empresa_id"], 'firma')
    if firma_path:
        insertar_imagen_en_docx(doc, firma_path, width_cm=4.0)
    sello_path = await obtener_firma(contrato["empresa_id"], 'sello')
    if sello_path:
        insertar_imagen_en_docx(doc, sello_path, width_cm=3.0)

    # Guardar
    output_dir = os.path.join(TEMPLATES_DIR, "output", "facturas")
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"factura_{numero_factura.replace('-', '_')}.docx")
    doc.save(path)

    # Registrar en DB
    async with connection() as conn:
        await conn.execute(
            """INSERT INTO pagos (contrato_id, concepto, monto, fecha_factura,
                                   estado, numero_factura, comprobante)
            VALUES ($1, $2, $3, $4, 'facturado', $5, $6)
            ON CONFLICT DO NOTHING""",
            contrato_id, concepto, monto, fechas.hoy(), numero_factura, path,
        )

    log.info(f"Factura generada: {path}")
    return path
